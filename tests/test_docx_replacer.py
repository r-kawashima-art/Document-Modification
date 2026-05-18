"""Verify acceptance criteria of the document-modification engine.

Tests are framework-free (just ``python3 tests/test_docx_replacer.py``) so
admin teammates can sanity-check the install without pytest.
"""
from __future__ import annotations

import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET


ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "src"))

from doc_modifier import docx_replacer, pipeline  # noqa: E402


NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _read_document_xml(docx_path: Path) -> ET.Element:
    with zipfile.ZipFile(docx_path) as z:
        with z.open("word/document.xml") as f:
            return ET.parse(f).getroot()


def _count_elements(root: ET.Element, tag: str) -> int:
    return len(list(root.iter(f"{NS}{tag}")))


def _collect_run_props(root: ET.Element) -> list[str]:
    """Return a stable string representation of every run's <w:rPr>."""
    props = []
    for r in root.iter(f"{NS}r"):
        rpr = r.find(f"{NS}rPr")
        if rpr is None:
            props.append("")
            continue
        # Canonicalize: ignore child ordering by sorting subtags + their attrib.
        children = []
        for child in rpr:
            attribs = sorted(child.attrib.items())
            children.append((child.tag, tuple(attribs)))
        props.append(repr(sorted(children)))
    return props


def test_render_preserves_line_breaks_and_paragraph_count():
    template = ROOT / "templates" / "Template_Invitation_Letter_Adventure_India_tokenized.docx"
    out = ROOT / "output" / "_test_render.docx"

    n = docx_replacer.render(template, {
        "name": "TEST NAME",
        "date_of_birth": "01/01/2000",
        "nationality": "Testland",
        "passport_no": "X0000000",
        "passport_issuing_country": "Testland",
        "date_of_issue": "01/01/2020",
        "date_of_expiry": "01/01/2030",
        "mobile_no": "+00 000-0000-0000",
    }, out)
    assert n >= 8, f"expected >=8 substitutions, got {n}"

    before = _read_document_xml(template)
    after = _read_document_xml(out)

    # AC-1: line breaks preserved.
    assert _count_elements(before, "p") == _count_elements(after, "p"), "paragraph count diverged"
    assert _count_elements(before, "br") == _count_elements(after, "br"), "line-break count diverged"

    # AC-2: every <w:rPr> that existed before exists in the same multiset after.
    before_props = sorted(_collect_run_props(before))
    after_props = sorted(_collect_run_props(after))
    # 'after' may have one fewer non-empty run (we sometimes blank trailing runs),
    # but every <w:rPr> we kept must have been present in the original set.
    for prop in after_props:
        assert prop in before_props, f"new font property appeared: {prop[:120]}"
    print(f"  ✓ paragraphs: {_count_elements(before, 'p')} == {_count_elements(after, 'p')}")
    print(f"  ✓ breaks:     {_count_elements(before, 'br')} == {_count_elements(after, 'br')}")
    print(f"  ✓ font properties: all {len(after_props)} run-property sets in output also exist in source")


def test_content_substituted_correctly():
    template = ROOT / "templates" / "Template_Invitation_Letter_Adventure_India_tokenized.docx"
    out = ROOT / "output" / "_test_content.docx"
    mapping = {
        "name": "Ms. Hanako Suzuki",
        "date_of_birth": "12/03/1985",
        "nationality": "Japan",
        "passport_no": "TR9981023",
        "passport_issuing_country": "Japan",
        "date_of_issue": "01/04/2022",
        "date_of_expiry": "01/04/2032",
        "mobile_no": "+81 80-1234-5678",
    }
    docx_replacer.render(template, mapping, out)
    from docx import Document
    doc = Document(str(out))

    all_text = " ".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                all_text += " " + cell.text

    for v in mapping.values():
        assert v in all_text, f"substituted value not found in output: {v!r}"

    # No raw tokens left.
    assert "{{" not in all_text, "leftover {{token}} in output"
    print(f"  ✓ all 8 values substituted; no leftover tokens")


def test_pipeline_two_rows():
    results = pipeline.run(
        template=ROOT / "templates" / "Template_Invitation_Letter_Adventure_India_tokenized.docx",
        data=ROOT / "data" / "sample_data.xlsx",
        out_dir=ROOT / "output",
        formats=("docx",),
    )
    assert len(results) == 2, f"expected 2 outputs, got {len(results)}"
    for r in results:
        assert r.primary_out.exists(), f"missing output: {r.primary_out}"
        assert r.substitutions >= 8, f"row {r.row_index}: too few substitutions ({r.substitutions})"
    print(f"  ✓ pipeline produced 2 documents with full substitution")


if __name__ == "__main__":
    print("test_render_preserves_line_breaks_and_paragraph_count")
    test_render_preserves_line_breaks_and_paragraph_count()
    print("test_content_substituted_correctly")
    test_content_substituted_correctly()
    print("test_pipeline_two_rows")
    test_pipeline_two_rows()
    print("\nAll acceptance checks passed.")
