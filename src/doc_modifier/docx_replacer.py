"""Run-aware ``{{token}}`` replacement for ``.docx`` templates.

Why this exists
---------------
Microsoft Word frequently splits a single visible string across multiple
``<w:r>`` (run) elements because of invisible spell-check, formatting,
or revision boundaries. A naïve ``paragraph.text = ...`` would collapse
every run into one and discard each run's ``<w:rPr>``, destroying the
template's fonts (フォント).

This module:
    1. Walks every paragraph in the body and inside every table cell.
    2. Concatenates the runs' visible text into a single buffer, while
       remembering each character's originating run.
    3. Finds every ``{{key}}`` match with a regex.
    4. For each match, writes the replacement into the *starting* run
       and blanks out the trailing runs that contributed to the match —
       this guarantees that every run we leave non-empty keeps its
       original ``<w:rPr>`` (and therefore its font, size, weight,
       color, etc.) intact.

Because we never mutate ``<w:p>`` (paragraph) or ``<w:br>`` (line
break) elements, the template's line-break layout (改行レイアウト)
is preserved bit-for-bit.
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable, Mapping

from docx import Document
from docx.document import Document as _DocumentT
from docx.table import _Cell
from docx.text.paragraph import Paragraph


TOKEN_RE = re.compile(r"\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}")


def _iter_paragraphs(doc: _DocumentT) -> Iterable[Paragraph]:
    """Yield every paragraph in the document body *and* in every table cell.

    Headers and footers are also covered.
    """
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_cell_paragraphs(cell)
    for section in doc.sections:
        for hf in (section.header, section.footer):
            yield from hf.paragraphs
            for table in hf.tables:
                for row in table.rows:
                    for cell in row.cells:
                        yield from _iter_cell_paragraphs(cell)


def _iter_cell_paragraphs(cell: _Cell) -> Iterable[Paragraph]:
    yield from cell.paragraphs
    for table in cell.tables:
        for row in table.rows:
            for sub in row.cells:
                yield from _iter_cell_paragraphs(sub)


def _replace_in_paragraph(paragraph: Paragraph, mapping: Mapping[str, str]) -> int:
    """Replace every ``{{token}}`` in ``paragraph`` using ``mapping``.

    Returns the number of substitutions performed.
    """
    runs = paragraph.runs
    if not runs:
        return 0

    # Build a flat text buffer + a parallel array mapping buffer-index -> run-index.
    texts = [r.text or "" for r in runs]
    full = "".join(texts)
    if "{{" not in full:
        return 0

    # buffer index -> originating run index
    run_of: list[int] = []
    for ri, t in enumerate(texts):
        run_of.extend([ri] * len(t))

    # Collect substitutions left-to-right. We must apply them in reverse so
    # that earlier indices remain valid as we mutate text.
    matches: list[tuple[int, int, str]] = []  # (start, end, replacement)
    for m in TOKEN_RE.finditer(full):
        key = m.group(1)
        if key not in mapping:
            continue
        matches.append((m.start(), m.end(), mapping[key]))

    if not matches:
        return 0

    # Apply matches in reverse so positions stay valid.
    for start, end, repl in reversed(matches):
        start_run = run_of[start]
        end_run = run_of[end - 1]
        # local offsets inside the starting / ending runs
        # find where in the start_run the match begins
        run_start_in_buf = run_of.index(start_run)  # first char of that run
        local_start = start - run_start_in_buf
        # end_run end offset
        end_run_first = end - 1
        # walk back to find first index of end_run in buf
        while end_run_first > 0 and run_of[end_run_first - 1] == end_run:
            end_run_first -= 1
        local_end = end - end_run_first  # exclusive

        if start_run == end_run:
            run = runs[start_run]
            run.text = run.text[:local_start] + repl + run.text[local_end:]
        else:
            # Replacement lives in the starting run. Trailing chunk of the
            # ending run is preserved by appending to the start run's text
            # (its font wins — typically identical to the end run's anyway).
            start_run_text = runs[start_run].text
            end_run_text = runs[end_run].text
            runs[start_run].text = start_run_text[:local_start] + repl + end_run_text[local_end:]
            # Blank out everything strictly between start_run+1 .. end_run inclusive.
            for ri in range(start_run + 1, end_run + 1):
                runs[ri].text = ""

        # Rebuild run_of for the remaining (earlier) matches.
        texts = [r.text or "" for r in runs]
        run_of = []
        for ri, t in enumerate(texts):
            run_of.extend([ri] * len(t))

    return len(matches)


def render(template_path: str | Path, mapping: Mapping[str, str], out_path: str | Path) -> int:
    """Render ``template_path`` against ``mapping``, writing to ``out_path``.

    :returns: total number of token substitutions performed.
    """
    doc = Document(str(template_path))
    total = 0
    for p in _iter_paragraphs(doc):
        total += _replace_in_paragraph(p, mapping)
    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return total


def find_tokens(template_path: str | Path) -> set[str]:
    """Return every distinct token (e.g. ``'name'``) referenced in the template."""
    doc = Document(str(template_path))
    found: set[str] = set()
    for p in _iter_paragraphs(doc):
        for m in TOKEN_RE.finditer(p.text or ""):
            found.add(m.group(1))
    return found
