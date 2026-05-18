"""Auto-tokenize a source document into a template with ``{{snake_case}}`` placeholders.

Goal: remove the manual find-and-replace step from the template-onboarding
workflow described in README §5.

Two modes, combined by default:

1. **Auto-detect** — scan tables; for every row matching the *label : value*
   pattern (column 0 short label, last column value), generate a snake_case
   token from the label and stage a replacement.
2. **Explicit mapping** — user-supplied ``original_text → token_name`` map.
   Wins if it conflicts with auto-detect.

After table tokens are decided, we sweep the body for inline occurrences of
the table values (e.g. "Mr. Takamichi Yanai" also appearing in the body
paragraph) and reuse the same token. *Ambiguous* values that map to more
than one token (e.g. two table rows both containing "Japan") are kept
**cell-scoped** — applied only to their specific table cell, never swept
into the body, with a warning printed for review.

The actual text mutation uses the same run-aware logic as ``docx_replacer``
so fonts (フォント) and line breaks (改行) survive unchanged.
"""
from __future__ import annotations

import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable, Mapping

from docx import Document
from docx.document import Document as _DocumentT
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from openpyxl import Workbook, load_workbook

from .docx_replacer import TOKEN_RE  # re-used for verification
from .xlsx_loader import normalize_header

# Heuristic limits for "looks like a label : value table"
_LABEL_MAX_CHARS = 60
_VALUE_MAX_CHARS = 200


@dataclass
class Replacement:
    """One staged change.

    If ``cell_addr`` is set (table_idx, row_idx, col_idx), the replacement is
    applied **only** inside that cell — used for ambiguous values where the
    same literal text is associated with multiple tokens. Otherwise the
    replacement is applied document-wide.
    """
    original: str
    token: str
    label: str
    location: str
    cell_addr: tuple[int, int, int] | None = None  # (table, row, col)
    example_value: str = ""


@dataclass
class TokenizationPlan:
    replacements: list[Replacement] = field(default_factory=list)
    skipped: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)

    @property
    def tokens(self) -> list[str]:
        """Unique tokens in first-occurrence order."""
        seen: list[str] = []
        for r in self.replacements:
            if r.token not in seen:
                seen.append(r.token)
        return seen

    @property
    def token_to_example(self) -> dict[str, str]:
        out: dict[str, str] = {}
        for r in self.replacements:
            if r.token not in out:
                out[r.token] = r.example_value or r.original
        return out


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _slug(label: str) -> str:
    return normalize_header(label) or "field"


def _iter_body_paragraphs(doc: _DocumentT) -> Iterable[tuple[Paragraph, str]]:
    for i, p in enumerate(doc.paragraphs):
        yield p, f"body paragraph {i}"


def _iter_cell_paragraphs(cell: _Cell) -> Iterable[Paragraph]:
    yield from cell.paragraphs
    for t in cell.tables:
        for row in t.rows:
            for sub in row.cells:
                yield from _iter_cell_paragraphs(sub)


def _table_is_label_value(table: Table) -> bool:
    rows = list(table.rows)
    if len(rows) < 2:
        return False
    label_count = 0
    for row in rows:
        if len(row.cells) < 2:
            return False
        label = row.cells[0].text.strip()
        value = row.cells[-1].text.strip()
        if not label or not value:
            return False
        if len(label) > _LABEL_MAX_CHARS or len(value) > _VALUE_MAX_CHARS:
            return False
        if "." in label and not label.endswith("."):
            return False
        label_count += 1
    return label_count >= 2


# ---------------------------------------------------------------------------
# Run-aware text replacement (reused at apply-time)
# ---------------------------------------------------------------------------

def _replace_text_in_paragraph(paragraph: Paragraph, target: str, replacement: str) -> int:
    """Run-aware replace of ``target`` with ``replacement`` in one paragraph.

    Preserves formatting via the same algorithm as
    ``docx_replacer._replace_in_paragraph``. Returns the substitution count.
    """
    runs = paragraph.runs
    if not runs or not target:
        return 0

    texts = [r.text or "" for r in runs]
    full = "".join(texts)
    if target not in full:
        return 0

    run_of: list[int] = []
    for ri, t in enumerate(texts):
        run_of.extend([ri] * len(t))

    count = 0
    pos = full.rfind(target)
    while pos != -1:
        start, end = pos, pos + len(target)
        start_run = run_of[start]
        end_run = run_of[end - 1]

        run_start_first = start
        while run_start_first > 0 and run_of[run_start_first - 1] == start_run:
            run_start_first -= 1
        local_start = start - run_start_first

        end_run_first = end - 1
        while end_run_first > 0 and run_of[end_run_first - 1] == end_run:
            end_run_first -= 1
        local_end = end - end_run_first

        if start_run == end_run:
            r = runs[start_run]
            r.text = r.text[:local_start] + replacement + r.text[local_end:]
        else:
            runs[start_run].text = runs[start_run].text[:local_start] + replacement + runs[end_run].text[local_end:]
            for ri in range(start_run + 1, end_run + 1):
                runs[ri].text = ""

        count += 1
        texts = [r.text or "" for r in runs]
        full = "".join(texts)
        run_of = []
        for ri, t in enumerate(texts):
            run_of.extend([ri] * len(t))
        pos = full.rfind(target, 0, start)
    return count


def _replace_in_cell(cell: _Cell, target: str, replacement: str) -> int:
    count = 0
    for p in _iter_cell_paragraphs(cell):
        count += _replace_text_in_paragraph(p, target, replacement)
    return count


def _replace_doc_wide(doc: _DocumentT, target: str, replacement: str) -> int:
    count = 0
    for p in doc.paragraphs:
        count += _replace_text_in_paragraph(p, target, replacement)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                count += _replace_in_cell(cell, target, replacement)
    for section in doc.sections:
        for hf in (section.header, section.footer):
            for p in hf.paragraphs:
                count += _replace_text_in_paragraph(p, target, replacement)
            for table in hf.tables:
                for row in table.rows:
                    for cell in row.cells:
                        count += _replace_in_cell(cell, target, replacement)
    return count


# ---------------------------------------------------------------------------
# Mapping loaders
# ---------------------------------------------------------------------------

def load_mapping(path: str | Path) -> dict[str, str]:
    """Load an explicit ``original_text → token`` mapping from yaml/json/xlsx."""
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix == ".json":
        return {str(k): str(v) for k, v in json.loads(p.read_text("utf-8")).items()}
    if suffix in {".yaml", ".yml"}:
        try:
            import yaml  # type: ignore
        except ImportError as exc:
            raise RuntimeError("PyYAML is required for YAML mappings — `pip install pyyaml`.") from exc
        data = yaml.safe_load(p.read_text("utf-8"))
        return {str(k): str(v) for k, v in data.items()}
    if suffix == ".xlsx":
        wb = load_workbook(filename=str(p), data_only=True)
        ws = wb.active
        out: dict[str, str] = {}
        rows_iter = ws.iter_rows(values_only=True)
        next(rows_iter, None)
        for raw in rows_iter:
            if not raw or len(raw) < 2 or raw[0] is None:
                continue
            out[str(raw[0])] = str(raw[1])
        return out
    raise ValueError(f"Unsupported mapping file type: {suffix}")


# ---------------------------------------------------------------------------
# analyze / apply / emit_starter_data
# ---------------------------------------------------------------------------

def analyze(
    source_path: str | Path,
    explicit_mapping: Mapping[str, str] | None = None,
    auto_detect_tables: bool = True,
) -> TokenizationPlan:
    """Build a :class:`TokenizationPlan` without modifying the source.

    Table rows are recorded as **cell-scoped** replacements so two rows with
    the same value (e.g. both "Japan") can still receive distinct tokens.
    A body sweep then promotes *unambiguous* values (single token per value)
    to document-wide replacements; ambiguous values stay cell-scoped and a
    warning is recorded.
    """
    doc = Document(str(source_path))
    plan = TokenizationPlan()

    # ---- 1) auto-detect tables (cell-scoped) -----------------------------
    if auto_detect_tables:
        for ti, table in enumerate(doc.tables):
            if not _table_is_label_value(table):
                plan.skipped.append(f"table[{ti}] does not look like label:value")
                continue
            for ri, row in enumerate(table.rows):
                label = row.cells[0].text.strip()
                value = row.cells[-1].text.strip()
                if not value:
                    continue
                last_col = len(row.cells) - 1
                token = _slug(label)
                plan.replacements.append(Replacement(
                    original=value,
                    token=token,
                    label=label,
                    location=f"table[{ti}].row[{ri}].col[{last_col}]",
                    cell_addr=(ti, ri, last_col),
                    example_value=value,
                ))

    # ---- 2) explicit overrides (doc-wide unless overlapping a table) ------
    if explicit_mapping:
        for original, token in explicit_mapping.items():
            token = re.sub(r"^\{\{|\}\}$", "", token).strip()
            plan.replacements.append(Replacement(
                original=original,
                token=token,
                label=token.replace("_", " ").title(),
                location="(explicit mapping)",
                cell_addr=None,
                example_value=original,
            ))

    # ---- 3) decide which originals are safe for doc-wide body sweep -------
    # An original is "safe" only if every replacement bearing that original
    # text agrees on a single token. Otherwise it's ambiguous and we keep all
    # of its replacements cell-scoped.
    original_to_tokens: dict[str, set[str]] = {}
    for r in plan.replacements:
        original_to_tokens.setdefault(r.original, set()).add(r.token)

    safe_originals: dict[str, str] = {
        orig: tokens.pop() for orig, tokens in original_to_tokens.items() if len(tokens) == 1
    }
    ambiguous = {orig for orig, tokens in original_to_tokens.items() if len(tokens) > 1}
    for orig in ambiguous:
        plan.warnings.append(
            f"Ambiguous value {orig!r} maps to multiple tokens "
            f"({', '.join(sorted(original_to_tokens[orig]))}); keeping cell-scoped only."
        )

    # ---- 4) sweep body for inline duplicates of *safe* values -------------
    for orig, tok in safe_originals.items():
        for p, hint in _iter_body_paragraphs(doc):
            if orig in (p.text or "") and not any(
                x.original == orig and x.location == hint for x in plan.replacements
            ):
                plan.replacements.append(Replacement(
                    original=orig,
                    token=tok,
                    label=orig,
                    location=hint,
                    cell_addr=None,
                    example_value=orig,
                ))

    return plan


def apply(source_path: str | Path, plan: TokenizationPlan, out_path: str | Path) -> int:
    """Apply ``plan`` and write to ``out_path``. Returns total substitutions."""
    doc = Document(str(source_path))
    total = 0

    # Apply cell-scoped first (longest original first so substrings don't shadow).
    cell_scoped = sorted(
        [r for r in plan.replacements if r.cell_addr is not None],
        key=lambda r: -len(r.original),
    )
    for r in cell_scoped:
        ti, ri, ci = r.cell_addr  # type: ignore[misc]
        try:
            cell = doc.tables[ti].rows[ri].cells[ci]
        except IndexError:
            plan.warnings.append(f"cell {r.cell_addr} not found at apply time; skipped {r.token}")
            continue
        total += _replace_in_cell(cell, r.original, "{{" + r.token + "}}")

    # Then doc-wide replacements (e.g. inline body matches, explicit mappings).
    doc_wide = sorted(
        [r for r in plan.replacements if r.cell_addr is None],
        key=lambda r: -len(r.original),
    )
    seen_targets: set[str] = set()
    for r in doc_wide:
        if r.original in seen_targets:
            continue
        seen_targets.add(r.original)
        total += _replace_doc_wide(doc, r.original, "{{" + r.token + "}}")

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return total


def emit_starter_data(plan: TokenizationPlan, out_path: str | Path) -> Path:
    """Emit a starter .xlsx data file with token columns + one example row."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Data"

    tokens = plan.tokens or ["name"]
    ws.append([*tokens, "output_filename"])
    example_row = [plan.token_to_example.get(t, "") for t in tokens]
    example_row.append("output_001")
    ws.append(example_row)

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(out_path))
    return out_path


def verify(out_docx_path: str | Path) -> set[str]:
    doc = Document(str(out_docx_path))
    found: set[str] = set()

    def scan(paragraphs):
        for p in paragraphs:
            for m in TOKEN_RE.finditer(p.text or ""):
                found.add(m.group(1))

    scan(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                scan(cell.paragraphs)
    return found
